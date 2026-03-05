"""
Multi-format Document Loader
Handles PDF, DOCX, TXT, HTML, and Image files for RAG ingestion.
"""
import logging
from pathlib import Path

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Unified document loader supporting multiple file formats."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".html", ".htm", ".png", ".jpg", ".jpeg"}

    def __init__(self):
        pass

    def load(self, file_path: str | Path) -> list[Document]:
        """Load a document and return a list of LangChain Document objects."""
        file_path = Path(file_path)
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return []

        ext = file_path.suffix.lower()
        loader_map = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".txt": self._load_text,
            ".html": self._load_html,
            ".htm": self._load_html,
            ".png": self._load_image,
            ".jpg": self._load_image,
            ".jpeg": self._load_image,
        }

        loader_fn = loader_map.get(ext)
        if loader_fn is None:
            logger.warning(f"Unsupported file type: {ext}")
            return []

        try:
            docs = loader_fn(file_path)
            logger.info(f"Loaded {len(docs)} document(s) from {file_path.name}")
            return docs
        except Exception as e:
            logger.error(f"Error loading {file_path.name}: {e}")
            return []

    def _load_pdf(self, file_path: Path) -> list[Document]:
        """Load PDF using PyMuPDF."""
        import fitz  # PyMuPDF

        docs = []
        pdf = fitz.open(str(file_path))
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text = page.get_text("text")
            if text.strip():
                docs.append(Document(
                    page_content=text.strip(),
                    metadata={
                        "source": file_path.name,
                        "source_type": "uploaded_pdf",
                        "page": page_num + 1,
                        "total_pages": len(pdf),
                    }
                ))
        pdf.close()
        return docs

    def _load_docx(self, file_path: Path) -> list[Document]:
        """Load DOCX using python-docx."""
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        if full_text:
            return [Document(
                page_content="\n\n".join(full_text),
                metadata={
                    "source": file_path.name,
                    "source_type": "uploaded_docx",
                    "paragraphs": len(full_text),
                }
            )]
        return []

    def _load_text(self, file_path: Path) -> list[Document]:
        """Load plain text files."""
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        if text.strip():
            return [Document(
                page_content=text.strip(),
                metadata={
                    "source": file_path.name,
                    "source_type": "uploaded_txt",
                }
            )]
        return []

    def _load_html(self, file_path: Path) -> list[Document]:
        """Load and clean HTML files."""
        import html2text
        from bs4 import BeautifulSoup

        raw = file_path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(raw, "lxml")

        # Remove script/style tags
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()

        converter = html2text.HTML2Text()
        converter.ignore_links = False
        converter.ignore_images = True
        converter.body_width = 0
        text = converter.handle(str(soup))

        if text.strip():
            title = soup.title.get_text(strip=True) if soup.title else file_path.name
            return [Document(
                page_content=text.strip(),
                metadata={
                    "source": file_path.name,
                    "source_type": "uploaded_html",
                    "title": title,
                }
            )]
        return []

    def _load_image(self, file_path: Path) -> list[Document]:
        """Load image and extract text via OCR (Tesseract)."""
        try:
            from PIL import Image
            import pytesseract

            img = Image.open(str(file_path))
            # Try both English and Arabic OCR
            text = pytesseract.image_to_string(img, lang="eng+ara")
            if text.strip():
                return [Document(
                    page_content=text.strip(),
                    metadata={
                        "source": file_path.name,
                        "source_type": "uploaded_image",
                        "ocr_engine": "tesseract",
                    }
                )]
        except ImportError:
            logger.warning("Tesseract not available. Image OCR skipped.")
        except Exception as e:
            logger.warning(f"OCR failed for {file_path.name}: {e}")

        return []


def load_documents(file_paths: list[str | Path]) -> list[Document]:
    """Convenience function to load multiple documents."""
    loader = DocumentLoader()
    all_docs = []
    for fp in file_paths:
        all_docs.extend(loader.load(fp))
    return all_docs
